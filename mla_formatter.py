"""
mla_formatter.py
Core engine for one-click MLA formatting (MLA Handbook, 9th Edition).

Two entry points:
  create_blank_template(...)      -> blank MLA-formatted docx, ready to type into
  convert_draft_to_mla(...)       -> takes an existing .docx or .txt draft and
                                      reformats it into proper MLA layout
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- optional heuristic classifiers (opt-in only) ----------

def is_likely_heading(text):
    """Detect section headings in body text (e.g. Introduction, 1. Background).
    Exposed for GUI opt-in; NOT used by default."""
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) > 80:
        return False
    if stripped[-1] in ".!?":
        return False
    if len(stripped) >= 3 and stripped.isupper():
        return True
    if re.match(r'^[IVXivx]+\.', stripped):
        return True
    if re.match(r'^\d+(\.\d+)*\.?\s', stripped):
        return True
    words = stripped.split()
    if 2 <= len(words) <= 8:
        upper_words = sum(1 for w in words if w and w[0].isupper())
        if upper_words >= len(words) * 0.6:
            return True
    if len(words) == 1 and stripped[0].isupper() and len(stripped) > 3:
        return True
    return False


def is_likely_block_quote(text):
    """Check if text starts with > (Markdown block quote marker).
    Returns True only on explicit > prefix; length-based detection removed."""
    stripped = text.strip()
    return stripped.startswith(">")


# ---------- low-level MLA building blocks ----------

def set_mla_page_setup(doc):
    """8.5 x 11 inch, 1-inch margins (MLA 9th para 1.1)."""
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_default_font(doc, font_name="Times New Roman", size=12):
    """Times New Roman 12 pt for all scripts (MLA 9th para 1.2)."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font_name)


def _set_update_fields(doc):
    """Ensure Word auto-updates PAGE fields when opened (MLA 9th para 1.5)."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _extract_last_name(name, override=None):
    """Return header last name: override if given, otherwise last word of name."""
    if override and override.strip():
        return override.strip()
    if name.strip():
        return name.strip().split()[-1]
    return "LastName"


def add_header_with_pagenum(doc, last_name):
    """LastName + auto PAGE field, top-right (MLA 9th para 1.5)."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{last_name} ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # PAGE field: three separate runs per OOXML spec
    for _tag, _text in [("begin", None), ("instrText", " PAGE "), ("end", None)]:
        r = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), _tag)
        r._r.append(fld)
        if _text:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = _text
            r._r.append(instr)


def set_double_spacing(paragraph):
    """Double spacing, 0 pt before/after (MLA 9th para 1.3-1.4)."""
    pf = paragraph.paragraph_format
    pf.line_spacing = 2.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def add_mla_heading_block(doc, name, instructor, course, date_str, title):
    """First-page heading block at top left + centered title (MLA 9th para 2)."""
    for text in (name, instructor, course, date_str):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_double_spacing(p)
    # Title: centered, no bold, no underline, same font/size
    title_p = doc.add_paragraph(title)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_double_spacing(title_p)


def add_body_paragraph(doc, text, works_cited_entry=False,
                       is_heading=False, is_block_quote=False):
    """Add a paragraph with MLA-formatted indentation and spacing.

    Parameters
    ----------
    works_cited_entry : bool
        Hanging indent (left +0.5 in, first-line -0.5 in). MLA 9th para 8.
    is_heading : bool
        Left-aligned, bold, no indent. (Only when user opts in.)
    is_block_quote : bool
        Indented 0.5 in from left, no first-line indent. MLA 9th para 6.35.
        (Prose quotations >4 lines.)
    """
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_double_spacing(p)

    if works_cited_entry:
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    elif is_block_quote:
        p.paragraph_format.left_indent = Inches(0.5)   # MLA 9th: 0.5 in
        p.paragraph_format.first_line_indent = Inches(0)
    elif is_heading:
        p.paragraph_format.first_line_indent = Inches(0)
        for run in p.runs:
            run.bold = True
    else:
        p.paragraph_format.first_line_indent = Inches(0.5)

    return p


# ---------- Works Cited helpers ----------

_WORKS_CITED_RE = re.compile(
    r"^(?:Works?\s+Cited|Work\s+Cited)[.\-:;!?]*$", re.IGNORECASE
)


def _wc_sort_key(entry):
    """Sort key for Works Cited entries: first non-article word, case-insensitive."""
    text = entry.strip()
    # Strip leading articles per MLA convention (A, An, The)
    base = re.sub(r'^(A|An|The)\s+', '', text, flags=re.IGNORECASE)
    # Strip leading non-word characters for sort key
    return re.sub(r'^[\W_]+', '', base.lower()).strip()


# ---------- validation / helpers ----------

def validate_inputs(name, title):
    errors = []
    if not name.strip():
        errors.append("Student Name cannot be empty.")
    if not title.strip():
        errors.append("Essay Title cannot be empty.")
    return errors


def generate_filename(name, last_name_override=None):
    """Auto-generate filename: LastName_MLA_Essay.docx"""
    ln = _extract_last_name(name, last_name_override)
    safe = re.sub(r'[\\/*?:"<>|]', "", ln) or "MLA"
    return f"{safe}_MLA_Essay.docx"


def format_summary(name, last_name_override=None):
    """Return a human-readable list of MLA formatting specs."""
    ln = _extract_last_name(name, last_name_override)
    return [
        f"Header: {ln} 1 (upper right)",
        "Paper: 8.5 x 11 inch (Letter)",
        "Font: Times New Roman 12 pt",
        "Margins: 1 inch (all sides)",
        "Spacing: Double (0 pt before/after)",
        "First-line indent: 0.5 inch",
        "Alignment: Left (not justify)",
        "Block quote indent: 0.5 inch",
        "Works Cited: new page, centered, alpha order",
        "Works Cited entries: hanging indent 0.5 inch",
    ]


MLG_DISCLAIMER = (
    "This tool reformats plain essay text into MLA format. "
    "Complex formatting such as images, footnotes, tables, lists, "
    "and special styles may not be preserved."
)


# ---------- entry point 1: blank template ----------

def create_blank_template(output_path, name="Your Name", instructor="Instructor Name",
                           course="Course Number", date_str="Day Month Year",
                           title="Title of Your Paper", last_name_override=None):
    doc = Document()
    set_mla_page_setup(doc)
    set_default_font(doc)
    _set_update_fields(doc)
    last_name = _extract_last_name(name, last_name_override)
    add_header_with_pagenum(doc, last_name)
    add_mla_heading_block(doc, name, instructor, course, date_str, title)
    add_body_paragraph(doc, "[Start writing your essay here.]")
    doc.save(output_path)


# ---------- entry point 2: convert an existing draft ----------

def _extract_paragraphs(input_path, txt_paragraph_mode="blank_line"):
    """Extract paragraphs from a .docx or .txt file.

    Parameters
    ----------
    txt_paragraph_mode : "blank_line" | "line_by_line"
        Only applies to .txt input.
    """
    if input_path.lower().endswith(".docx"):
        src = Document(input_path)
        return [p.text for p in src.paragraphs if p.text.strip()]

    # .txt
    with open(input_path, encoding="utf-8-sig") as f:
        text = f.read()
    normalized = text.replace("\r\n", "\n")

    if txt_paragraph_mode == "line_by_line":
        return [line.strip() for line in normalized.split("\n") if line.strip()]
    # default: blank_line mode
    return [chunk.strip() for chunk in re.split(r"\n\s*\n", normalized) if chunk.strip()]


def convert_draft_to_mla(input_path, output_path, name, instructor, course,
                          date_str, title, last_name_override=None,
                          txt_paragraph_mode="blank_line",
                          enable_heading_detection=False,
                          enable_block_quote=False):
    """Convert a draft .docx/.txt into an MLA-formatted .docx (9th Edition).

    Parameters
    ----------
    last_name_override : str | None
        Manual override for the page header last name.
    txt_paragraph_mode : "blank_line" | "line_by_line"
        How .txt files are split into paragraphs.
    enable_heading_detection : bool
        If True, apply is_likely_heading heuristics for bold headings.
    enable_block_quote : bool
        If True, treat >-prefixed lines as block quotes (0.5 in indent).
    """
    paragraphs = _extract_paragraphs(input_path, txt_paragraph_mode)

    doc = Document()
    set_mla_page_setup(doc)
    set_default_font(doc)
    _set_update_fields(doc)
    last_name = _extract_last_name(name, last_name_override)
    add_header_with_pagenum(doc, last_name)
    add_mla_heading_block(doc, name, instructor, course, date_str, title)

    # -- Phase 1: write body paragraphs, buffer Works Cited entries --
    in_works_cited = False
    wc_entries = []

    for para_text in paragraphs:
        cleaned = para_text.strip()

        # --- Detect "Works Cited" heading ---
        if _WORKS_CITED_RE.match(cleaned):
            in_works_cited = True
            continue

        # --- Works Cited entries: buffer for sorting ---
        if in_works_cited:
            wc_entries.append(para_text)
            continue

        # --- Opt-in heading detection ---
        if enable_heading_detection and is_likely_heading(cleaned):
            add_body_paragraph(doc, para_text, is_heading=True)
            continue

        # --- Opt-in block quote (starts with >) ---
        if enable_block_quote and is_likely_block_quote(cleaned):
            display_text = cleaned.lstrip(">").strip()
            add_body_paragraph(doc, display_text, is_block_quote=True)
            continue

        # --- Normal body paragraph ---
        add_body_paragraph(doc, para_text)

    # -- Phase 2: Works Cited page (MLA 9th para 8) --
    if wc_entries:
        # Page break -> new page
        run = doc.add_paragraph().add_run()
        run.add_break(WD_BREAK.PAGE)

        # Works Cited heading: centered, no extra formatting
        wc_title = doc.add_paragraph("Works Cited")
        wc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_double_spacing(wc_title)

        # Sort alphabetically by first meaningful word
        for entry in sorted(wc_entries, key=_wc_sort_key):
            add_body_paragraph(doc, entry, works_cited_entry=True)

    doc.save(output_path)


if __name__ == "__main__":
    import tempfile
    tmp = tempfile.gettempdir()

    print("--- Self-test ---")
    create_blank_template(os.path.join(tmp, "test_blank.docx"),
                           name="Your Name", instructor="Instructor Name",
                           course="Course Number", date_str="20 June 2026",
                           title="Title of Your Paper")

    sample = ("First paragraph.\n\n"
              "Second paragraph.\n\n"
              "Works Cited\n\n"
              "Doe, Jane. Another Title. Publisher, 2021.\n\n"
              "Smith, John. Some Book Title. Publisher, 2020.\n\n"
              "Brown, Alice. Third Source. Publisher, 2019.")
    with open(os.path.join(tmp, "test_draft.txt"), "w", encoding="utf-8") as f:
        f.write(sample)

    convert_draft_to_mla(os.path.join(tmp, "test_draft.txt"),
                          os.path.join(tmp, "test_converted.docx"),
                          name="Your Name", instructor="Instructor Name",
                          course="Course Number", date_str="20 June 2026",
                          title="Title of Your Paper")
    print("Self-test OK.")
